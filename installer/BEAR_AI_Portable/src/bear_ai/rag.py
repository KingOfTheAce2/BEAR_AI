from __future__ import annotations
import os
import re
import json
import hashlib
from dataclasses import dataclass, asdict
from typing import List, Tuple, Dict, Any, Optional
from pathlib import Path
import logging
from collections import Counter
import time


_WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9_']+")

# Logging setup
logger = logging.getLogger(__name__)

# Document chunk settings for large legal documents
MAX_CHUNK_SIZE = 1000  # Maximum characters per chunk
CHUNK_OVERLAP = 100    # Overlap between chunks
MIN_CHUNK_SIZE = 50    # Minimum chunk size to keep


def _tokenize(text: str) -> List[str]:
    return [w.lower() for w in _WORD_RE.findall(text)]


def _read_text_file(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception:
        return ""


def _extract_pdf_text(path: str) -> str:
    """Extract text from PDF files with error handling and memory efficiency."""
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        logger.warning("pypdf not installed, PDF support disabled")
        return ""
    
    try:
        reader = PdfReader(path)
        parts: List[str] = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"[Page {i+1}] {text}")
                # Memory management for large PDFs
                if len(parts) > 100 and len('\n'.join(parts)) > 1000000:  # 1MB text limit
                    logger.warning(f"Large PDF detected ({path}), truncating at page {i+1}")
                    break
            except Exception as e:
                logger.error(f"Error extracting page {i+1} from {path}: {e}")
                continue
        return "\n\n".join(parts)
    except Exception as e:
        logger.error(f"Error reading PDF {path}: {e}")
        return ""


def _extract_docx_text(path: str) -> str:
    """Extract text from DOCX files with enhanced structure preservation."""
    try:
        import docx  # type: ignore
    except ImportError:
        logger.warning("python-docx not installed, DOCX support disabled")
        return ""
    
    try:
        doc = docx.Document(path)
        parts: List[str] = []
        
        # Extract paragraphs with basic structure
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure if available
                if para.style.name.startswith('Heading'):
                    parts.append(f"\n## {text}\n")
                else:
                    parts.append(text)
        
        # Extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)
            if table_text:
                parts.append("\n[Table]\n" + "\n".join(table_text) + "\n")
        
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"Error reading DOCX {path}: {e}")
        return ""


def _extract_markdown_text(path: str) -> str:
    """Extract and lightly process Markdown files."""
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
            # Basic markdown processing - remove excessive formatting
            content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)  # Remove heading markers
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Remove bold
            content = re.sub(r'\*(.*?)\*', r'\1', content)  # Remove italic
            content = re.sub(r'`(.*?)`', r'\1', content)  # Remove inline code
            content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)  # Remove links, keep text
            return content
    except Exception as e:
        logger.error(f"Error reading Markdown {path}: {e}")
        return ""


def _chunk_text(text: str, max_size: int = MAX_CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Smart chunking for large documents with overlap for context preservation."""
    if len(text) <= max_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + max_size
        
        if end >= len(text):
            # Last chunk
            chunk = text[start:]
        else:
            # Find good break points (sentence endings, paragraph breaks)
            chunk_text = text[start:end]
            # Look for sentence endings near the end
            for break_char in ['.\n\n', '.\n', '. ', '!\n', '?\n']:
                break_pos = chunk_text.rfind(break_char)
                if break_pos > max_size * 0.7:  # At least 70% of max size
                    end = start + break_pos + len(break_char)
                    break
            chunk = text[start:end]
        
        if len(chunk.strip()) >= MIN_CHUNK_SIZE:
            chunks.append(chunk.strip())
        
        # Move start with overlap
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks


@dataclass
class DocumentMetadata:
    """Metadata for indexed documents."""
    file_path: str
    file_hash: str
    file_size: int
    modified_time: float
    chunk_count: int
    indexed_time: float
    file_type: str


@dataclass
class Snippet:
    file: str
    text: str
    score: float
    chunk_id: Optional[int] = None
    metadata: Optional[Dict[str, Any]] = None


class DocumentIndex:
    """Document indexing and caching system for efficient retrieval."""
    
    def __init__(self, case_id: str, base_dir: str):
        self.case_id = case_id
        self.base_dir = Path(base_dir)
        self.docs_dir = self.base_dir / "docs" / case_id
        self.index_dir = self.base_dir / "index" / case_id
        self.index_dir.mkdir(parents=True, exist_ok=True)
        self.metadata_file = self.index_dir / "metadata.json"
        self.index_file = self.index_dir / "index.json"
        
        self._metadata: Dict[str, DocumentMetadata] = {}
        self._index: Dict[str, List[Tuple[str, str, int]]] = {}  # file -> [(text, hash, chunk_id)]
        self._load_metadata()
        self._load_index()
    
    def _load_metadata(self):
        """Load document metadata from disk."""
        if self.metadata_file.exists():
            try:
                with open(self.metadata_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self._metadata = {
                        k: DocumentMetadata(**v) for k, v in data.items()
                    }
            except Exception as e:
                logger.error(f"Error loading metadata: {e}")
                self._metadata = {}
    
    def _save_metadata(self):
        """Save document metadata to disk."""
        try:
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump({
                    k: asdict(v) for k, v in self._metadata.items()
                }, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving metadata: {e}")
    
    def _load_index(self):
        """Load document index from disk."""
        if self.index_file.exists():
            try:
                with open(self.index_file, 'r', encoding='utf-8') as f:
                    self._index = json.load(f)
            except Exception as e:
                logger.error(f"Error loading index: {e}")
                self._index = {}
    
    def _save_index(self):
        """Save document index to disk."""
        try:
            with open(self.index_file, 'w', encoding='utf-8') as f:
                json.dump(self._index, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving index: {e}")
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Get MD5 hash of file for change detection."""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception:
            return ""
    
    def _needs_reindex(self, file_path: Path) -> bool:
        """Check if file needs to be reindexed."""
        file_key = str(file_path)
        if file_key not in self._metadata:
            return True
        
        try:
            stat = file_path.stat()
            meta = self._metadata[file_key]
            return (stat.st_mtime > meta.modified_time or 
                    stat.st_size != meta.file_size or
                    self._get_file_hash(file_path) != meta.file_hash)
        except Exception:
            return True
    
    def update_index(self, force_rebuild: bool = False):
        """Update index with new or modified documents."""
        if not self.docs_dir.exists():
            return
        
        updated_files = 0
        for file_path in self.docs_dir.iterdir():
            if not file_path.is_file():
                continue
                
            file_key = str(file_path)
            if not force_rebuild and not self._needs_reindex(file_path):
                continue
            
            try:
                # Extract text based on file type
                name_lower = file_path.name.lower()
                if name_lower.endswith('.txt'):
                    text = _read_text_file(str(file_path))
                    file_type = 'txt'
                elif name_lower.endswith('.pdf'):
                    text = _extract_pdf_text(str(file_path))
                    file_type = 'pdf'
                elif name_lower.endswith('.docx'):
                    text = _extract_docx_text(str(file_path))
                    file_type = 'docx'
                elif name_lower.endswith(('.md', '.markdown')):
                    text = _extract_markdown_text(str(file_path))
                    file_type = 'markdown'
                else:
                    continue
                
                if not text.strip():
                    continue
                
                # Chunk the document
                chunks = _chunk_text(text)
                
                # Update index
                file_hash = self._get_file_hash(file_path)
                stat = file_path.stat()
                
                self._metadata[file_key] = DocumentMetadata(
                    file_path=file_key,
                    file_hash=file_hash,
                    file_size=stat.st_size,
                    modified_time=stat.st_mtime,
                    chunk_count=len(chunks),
                    indexed_time=time.time(),
                    file_type=file_type
                )
                
                self._index[file_key] = [
                    (chunk, file_hash, i) for i, chunk in enumerate(chunks)
                ]
                
                updated_files += 1
                logger.info(f"Indexed {file_path.name}: {len(chunks)} chunks")
                
            except Exception as e:
                logger.error(f"Error indexing {file_path}: {e}")
        
        if updated_files > 0:
            self._save_metadata()
            self._save_index()
            logger.info(f"Updated index for {updated_files} files")
    
    def get_all_chunks(self) -> List[Tuple[str, str, int, str]]:
        """Get all indexed chunks: (file_path, text, chunk_id, file_hash)."""
        chunks = []
        for file_path, file_chunks in self._index.items():
            for text, file_hash, chunk_id in file_chunks:
                chunks.append((file_path, text, chunk_id, file_hash))
        return chunks


class RAGPipeline:
    """Enhanced RAG pipeline with multi-format support and smart chunking.

    Features:
    - Supports PDF, DOCX, TXT, and Markdown files
    - Smart document chunking for large legal documents
    - Document indexing and caching for performance
    - Memory-efficient processing
    - Incremental updates
    """

    def __init__(self, docs: List[Tuple[str, str]] | None = None, case_id: str | None = None, base_dir: str | None = None):
        # Backwards compatibility
        self.docs = docs or []
        self.case_id = case_id
        self.base_dir = base_dir
        self.index: Optional[DocumentIndex] = None
        
        if case_id and base_dir:
            self.index = DocumentIndex(case_id, base_dir)
            # Auto-update index on initialization
            self.index.update_index()

    @classmethod
    def from_case_docs(cls, case_id: str, base_dir: str) -> "RAGPipeline":
        """Create RAG pipeline from case documents with indexing."""
        return cls(case_id=case_id, base_dir=base_dir)
    
    def update_documents(self, force_rebuild: bool = False):
        """Update document index."""
        if self.index:
            self.index.update_index(force_rebuild=force_rebuild)
        else:
            logger.warning("No index available - use from_case_docs() to enable indexing")
    
    def _get_docs_for_query(self) -> List[Tuple[str, str]]:
        """Get documents for querying - either from index or legacy docs list."""
        if self.index:
            chunks = self.index.get_all_chunks()
            return [(os.path.basename(file_path), text) for file_path, text, _, _ in chunks]
        else:
            return [(os.path.basename(path), text) for path, text in self.docs]

    def query(self, query_text: str, top_k: int = 3, min_score: float = 0.01) -> List[Snippet]:
        """Enhanced query with better relevance scoring and snippet extraction."""
        q_tokens = _tokenize(query_text)
        if not q_tokens:
            return []
        
        docs = self._get_docs_for_query()
        if not docs:
            return []
        
        # Build document frequency for IDF calculation
        df = Counter()
        doc_tokens = []
        for _, txt in docs:
            toks = set(_tokenize(txt))
            doc_tokens.append(toks)
            df.update(toks)
        
        N = max(len(docs), 1)
        q_token_set = set(q_tokens)
        
        results: List[Snippet] = []
        for i, (path, txt) in enumerate(docs):
            toks = _tokenize(txt)
            if not toks:
                continue
            
            # Enhanced scoring: TF-IDF with query term frequency
            score = 0.0
            query_matches = 0
            
            # Count tokens for TF calculation
            token_counts = Counter(toks)
            
            for t in q_token_set:
                if t in token_counts:
                    tf = token_counts[t] / len(toks)
                    idf = max(0.1, 1.0 + (N / max(df[t], 1)))
                    # Boost score for exact query term matches
                    boost = q_tokens.count(t)  # How many times this term appears in query
                    score += tf * idf * boost
                    query_matches += 1
            
            # Bonus for documents matching multiple query terms
            if query_matches > 1:
                score *= (1.0 + 0.2 * (query_matches - 1))
            
            if score >= min_score:
                # Extract better snippet around query terms
                snippet = self._extract_snippet(txt, q_tokens, max_length=400)
                
                # Get chunk metadata if available
                chunk_id = None
                metadata = None
                if self.index:
                    chunk_info = self.index.get_all_chunks()
                    if i < len(chunk_info):
                        _, _, chunk_id, _ = chunk_info[i]
                        if path in [os.path.basename(fp) for fp, _, _, _ in chunk_info]:
                            metadata = {
                                'file_type': getattr(self.index._metadata.get(path), 'file_type', 'unknown'),
                                'chunk_count': getattr(self.index._metadata.get(path), 'chunk_count', 1)
                            }
                
                results.append(Snippet(
                    file=path, 
                    text=snippet, 
                    score=score,
                    chunk_id=chunk_id,
                    metadata=metadata
                ))
        
        # Sort by score and return top results
        results.sort(key=lambda s: s.score, reverse=True)
        return results[:top_k]
    
    def _extract_snippet(self, text: str, query_tokens: List[str], max_length: int = 400) -> str:
        """Extract the most relevant snippet containing query terms."""
        if len(text) <= max_length:
            return text.strip()
        
        # Find positions of query terms (case-insensitive)
        text_lower = text.lower()
        positions = []
        for token in query_tokens:
            token_lower = token.lower()
            start = 0
            while True:
                pos = text_lower.find(token_lower, start)
                if pos == -1:
                    break
                positions.append(pos)
                start = pos + 1
        
        if not positions:
            # No matches found, return beginning of text
            return text[:max_length].strip() + "..."
        
        # Find the best window that contains the most query terms
        best_start = 0
        best_count = 0
        
        for start_pos in positions:
            # Count query terms in a window starting from this position
            window_end = min(start_pos + max_length, len(text))
            window_start = max(0, start_pos - max_length // 4)  # Start a bit before first match
            
            count = sum(1 for pos in positions if window_start <= pos < window_end)
            if count > best_count:
                best_count = count
                best_start = window_start
        
        # Extract the snippet
        end_pos = min(best_start + max_length, len(text))
        snippet = text[best_start:end_pos].strip()
        
        # Clean up snippet boundaries
        if best_start > 0:
            snippet = "..." + snippet
        if end_pos < len(text):
            snippet = snippet + "..."
        
        return snippet
    
    def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about indexed documents."""
        if not self.index:
            return {
                'total_documents': len(self.docs),
                'total_chunks': len(self.docs),
                'indexed': False
            }
        
        chunks = self.index.get_all_chunks()
        file_types = Counter()
        total_size = 0
        
        for file_path, meta in self.index._metadata.items():
            file_types[meta.file_type] += 1
            total_size += meta.file_size
        
        return {
            'total_documents': len(self.index._metadata),
            'total_chunks': len(chunks),
            'file_types': dict(file_types),
            'total_size_bytes': total_size,
            'indexed': True,
            'index_path': str(self.index.index_dir)
        }


def create_case_docs_structure(case_id: str, base_dir: str) -> Path:
    """Create the document directory structure for a case."""
    docs_dir = Path(base_dir) / "docs" / case_id
    docs_dir.mkdir(parents=True, exist_ok=True)
    return docs_dir