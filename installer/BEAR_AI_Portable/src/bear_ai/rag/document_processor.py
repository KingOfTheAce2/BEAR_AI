"""
Document Processing System
Extract, chunk, and prepare documents for RAG
"""

import asyncio
import logging
import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import uuid

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types"""
    TEXT = "text"
    MARKDOWN = "markdown"
    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"
    CSV = "csv"
    JSON = "json"
    XML = "xml"
    RTF = "rtf"
    UNKNOWN = "unknown"


@dataclass
class ProcessedDocument:
    """Processed document with chunks"""
    id: str
    title: str
    content: str
    doc_type: DocumentType
    chunks: List['DocumentChunk'] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    source_path: Optional[str] = None
    processed_at: float = field(default_factory=lambda: __import__('time').time())


@dataclass
class DocumentChunk:
    """Individual chunk of a document"""
    id: str
    document_id: str
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def __post_init__(self):
        if not self.id:
            self.id = str(uuid.uuid4())


class TextSplitter:
    """Split text into chunks with overlap"""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separator: str = "\n\n",
        keep_separator: bool = False
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separator = separator
        self.keep_separator = keep_separator
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        if not text or len(text) <= self.chunk_size:
            return [text] if text else []
        
        # Split by separator first
        if self.separator in text:
            return self._split_by_separator(text)
        else:
            return self._split_by_length(text)
    
    def _split_by_separator(self, text: str) -> List[str]:
        """Split text by separator with overlap"""
        parts = text.split(self.separator)
        chunks = []
        current_chunk = ""
        
        for part in parts:
            # If adding this part exceeds chunk size, finalize current chunk
            if current_chunk and len(current_chunk) + len(part) + len(self.separator) > self.chunk_size:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text
                
                if overlap_text and self.keep_separator:
                    current_chunk += self.separator
            
            # Add separator if keeping it and chunk is not empty
            if current_chunk and self.keep_separator:
                current_chunk += self.separator
            
            current_chunk += part
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _split_by_length(self, text: str) -> List[str]:
        """Split text by character length with overlap"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # Don't split in the middle of a word
            if end < len(text):
                # Look for word boundary
                while end > start and not text[end].isspace():
                    end -= 1
                
                # If no space found, use original end
                if end == start:
                    end = start + self.chunk_size
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position considering overlap
            start = max(start + 1, end - self.chunk_overlap)
        
        return chunks
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of current chunk"""
        if len(text) <= self.chunk_overlap:
            return text
        
        overlap_start = len(text) - self.chunk_overlap
        overlap_text = text[overlap_start:]
        
        # Try to start at word boundary
        space_idx = overlap_text.find(' ')
        if space_idx != -1:
            overlap_text = overlap_text[space_idx + 1:]
        
        return overlap_text


class DocumentProcessor:
    """Process documents for RAG system"""
    
    def __init__(self):
        self.text_splitter = TextSplitter()
        
        # Document type detection patterns
        self.type_patterns = {
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOCX,
            '.csv': DocumentType.CSV,
            '.json': DocumentType.JSON,
            '.xml': DocumentType.XML,
            '.rtf': DocumentType.RTF
        }
        
        logger.info("DocumentProcessor initialized")
    
    async def process_document(
        self,
        content: Union[str, Path],
        title: Optional[str] = None,
        doc_type: Optional[DocumentType] = None,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process a document into chunks"""
        
        doc_id = str(uuid.uuid4())
        source_path = None
        
        # Handle file path input
        if isinstance(content, (str, Path)) and Path(content).exists():
            file_path = Path(content)
            source_path = str(file_path)
            title = title or file_path.stem
            
            # Detect document type from extension
            if not doc_type:
                doc_type = self._detect_document_type(file_path)
            
            # Extract text from file
            text_content = await self._extract_text_from_file(file_path, doc_type)
        else:
            # Handle string content
            text_content = str(content) if content else ""
            title = title or f"Document_{doc_id[:8]}"
            doc_type = doc_type or DocumentType.TEXT
        
        if not text_content:
            logger.warning(f"No content extracted from document: {title}")
            return ProcessedDocument(
                id=doc_id,
                title=title,
                content="",
                doc_type=doc_type,
                metadata=metadata or {},
                source_path=source_path
            )
        
        # Create text splitter with specified parameters
        splitter = TextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        
        # Split into chunks
        chunk_texts = splitter.split_text(text_content)
        
        # Create chunk objects
        chunks = []
        current_char = 0
        
        for i, chunk_text in enumerate(chunk_texts):
            start_char = current_char
            end_char = start_char + len(chunk_text)
            
            chunk_metadata = (metadata or {}).copy()
            chunk_metadata.update({
                'document_title': title,
                'document_type': doc_type.value,
                'chunk_index': i,
                'total_chunks': len(chunk_texts)
            })
            
            chunk = DocumentChunk(
                id=str(uuid.uuid4()),
                document_id=doc_id,
                content=chunk_text,
                chunk_index=i,
                start_char=start_char,
                end_char=end_char,
                metadata=chunk_metadata
            )
            
            chunks.append(chunk)
            current_char = end_char
        
        # Create processed document
        processed_doc = ProcessedDocument(
            id=doc_id,
            title=title,
            content=text_content,
            doc_type=doc_type,
            chunks=chunks,
            metadata=metadata or {},
            source_path=source_path
        )
        
        logger.info(f"Processed document '{title}' into {len(chunks)} chunks")
        return processed_doc
    
    def _detect_document_type(self, file_path: Path) -> DocumentType:
        """Detect document type from file extension"""
        extension = file_path.suffix.lower()
        return self.type_patterns.get(extension, DocumentType.UNKNOWN)
    
    async def _extract_text_from_file(self, file_path: Path, doc_type: DocumentType) -> str:
        """Extract text from different file types"""
        
        try:
            if doc_type == DocumentType.TEXT:
                return await self._extract_text(file_path)
            
            elif doc_type == DocumentType.MARKDOWN:
                return await self._extract_markdown(file_path)
            
            elif doc_type == DocumentType.HTML:
                return await self._extract_html(file_path)
            
            elif doc_type == DocumentType.PDF:
                return await self._extract_pdf(file_path)
            
            elif doc_type == DocumentType.DOCX:
                return await self._extract_docx(file_path)
            
            elif doc_type == DocumentType.CSV:
                return await self._extract_csv(file_path)
            
            elif doc_type == DocumentType.JSON:
                return await self._extract_json(file_path)
            
            elif doc_type == DocumentType.XML:
                return await self._extract_xml(file_path)
            
            elif doc_type == DocumentType.RTF:
                return await self._extract_rtf(file_path)
            
            else:
                # Default to text extraction
                return await self._extract_text(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract plain text"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except:
                    continue
            
            # If all else fails, read as binary and decode with errors
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
    
    async def _extract_markdown(self, file_path: Path) -> str:
        """Extract text from Markdown files"""
        content = await self._extract_text(file_path)
        
        # Simple markdown to text conversion
        # Remove headers
        content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)
        
        # Remove emphasis markers
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Bold
        content = re.sub(r'\*(.*?)\*', r'\1', content)      # Italic
        content = re.sub(r'`(.*?)`', r'\1', content)        # Inline code
        
        # Remove links
        content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
        
        # Remove code blocks
        content = re.sub(r'```[^`]*```', '', content, flags=re.DOTALL)
        
        return content.strip()
    
    async def _extract_html(self, file_path: Path) -> str:
        """Extract text from HTML files"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except ImportError:
            logger.warning("BeautifulSoup not available for HTML parsing")
            return await self._extract_text(file_path)
    
    async def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            import pypdf
            
            text = ""
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            
            return text.strip()
            
        except ImportError:
            logger.warning("pypdf not available for PDF parsing")
            return ""
    
    async def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except ImportError:
            logger.warning("python-docx not available for DOCX parsing")
            return ""
    
    async def _extract_csv(self, file_path: Path) -> str:
        """Extract text from CSV files"""
        try:
            import csv
            
            text_parts = []
            
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                
                # Get headers
                headers = next(reader, [])
                if headers:
                    text_parts.append("Headers: " + ", ".join(headers))
                
                # Get sample rows
                for i, row in enumerate(reader):
                    if i >= 10:  # Limit to first 10 rows for processing
                        break
                    text_parts.append(" | ".join(str(cell) for cell in row))
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error parsing CSV {file_path}: {e}")
            return ""
    
    async def _extract_json(self, file_path: Path) -> str:
        """Extract text from JSON files"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            return self._json_to_text(data)
            
        except Exception as e:
            logger.error(f"Error parsing JSON {file_path}: {e}")
            return ""
    
    def _json_to_text(self, obj: Any, prefix: str = "") -> str:
        """Convert JSON object to readable text"""
        if isinstance(obj, dict):
            text_parts = []
            for key, value in obj.items():
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{prefix}{key}:")
                    text_parts.append(self._json_to_text(value, prefix + "  "))
                else:
                    text_parts.append(f"{prefix}{key}: {value}")
            return "\n".join(text_parts)
        
        elif isinstance(obj, list):
            text_parts = []
            for i, item in enumerate(obj[:10]):  # Limit to first 10 items
                text_parts.append(f"{prefix}Item {i+1}:")
                text_parts.append(self._json_to_text(item, prefix + "  "))
            return "\n".join(text_parts)
        
        else:
            return str(obj)
    
    async def _extract_xml(self, file_path: Path) -> str:
        """Extract text from XML files"""
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract all text content
            text_parts = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_parts.append(elem.text.strip())
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error parsing XML {file_path}: {e}")
            return ""
    
    async def _extract_rtf(self, file_path: Path) -> str:
        """Extract text from RTF files"""
        # RTF parsing is complex, fall back to text extraction
        logger.warning("RTF parsing not implemented, falling back to text extraction")
        return await self._extract_text(file_path)


# Global document processor instance
_global_processor: Optional[DocumentProcessor] = None

def get_document_processor() -> DocumentProcessor:
    """Get the global document processor"""
    global _global_processor
    if _global_processor is None:
        _global_processor = DocumentProcessor()
    return _global_processor

async def process_document(
    content: Union[str, Path],
    title: Optional[str] = None,
    doc_type: Optional[DocumentType] = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    metadata: Optional[Dict[str, Any]] = None
) -> ProcessedDocument:
    """Process a document"""
    processor = get_document_processor()
    return await processor.process_document(
        content, title, doc_type, chunk_size, chunk_overlap, metadata
    )

async def extract_text(file_path: Path) -> str:
    """Extract text from a file"""
    processor = get_document_processor()
    doc_type = processor._detect_document_type(file_path)
    return await processor._extract_text_from_file(file_path, doc_type)