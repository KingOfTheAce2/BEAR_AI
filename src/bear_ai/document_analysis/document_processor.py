"""
Enhanced Document Processing System for BEAR AI
Local processing with OCR, encryption, and legal document specialization

Features:
- PDF, DOCX, TXT parsing with Tesseract OCR
- Document type classification
- Text preprocessing and cleaning
- Secure document handling
- Performance optimization
"""

import asyncio
import hashlib
import io
import logging
import mimetypes
import os
import tempfile
import time
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Union, Tuple, Callable
import uuid

# Core dependencies
try:
    import fitz  # PyMuPDF for PDF processing
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Document type classification"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    HTML = "html"
    RTF = "rtf"
    XLSX = "xlsx"
    XLS = "xls"
    CSV = "csv"
    IMAGE = "image"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    """Document processing status"""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    SKIPPED = "skipped"


@dataclass
class ProcessingConfig:
    """Configuration for document processing"""
    # OCR Configuration
    ocr_enabled: bool = True
    ocr_language: str = "eng"
    ocr_dpi: int = 300
    ocr_psm: int = 6  # Tesseract Page Segmentation Mode

    # Text Processing
    chunk_size: int = 1000
    chunk_overlap: int = 200
    preserve_formatting: bool = True
    clean_text: bool = True

    # Performance
    max_workers: int = 4
    timeout_seconds: int = 300
    memory_limit_mb: int = 1000

    # Security
    encrypt_storage: bool = True
    audit_enabled: bool = True
    temp_file_cleanup: bool = True

    # Legal Document Features
    extract_clauses: bool = True
    detect_entities: bool = True
    analyze_risk: bool = True


@dataclass
class DocumentMetadata:
    """Document metadata"""
    file_path: Path
    file_name: str
    file_size: int
    mime_type: str
    document_type: DocumentType
    creation_time: Optional[float] = None
    modification_time: Optional[float] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    language: str = "en"
    encoding: str = "utf-8"
    checksum: Optional[str] = None

    def __post_init__(self):
        if self.checksum is None:
            self.checksum = self._calculate_checksum()

    def _calculate_checksum(self) -> str:
        """Calculate file checksum"""
        try:
            with open(self.file_path, 'rb') as f:
                content = f.read()
                return hashlib.sha256(content).hexdigest()
        except Exception:
            return ""


@dataclass
class ProcessingResult:
    """Document processing result"""
    document_id: str
    metadata: DocumentMetadata
    extracted_text: str
    cleaned_text: str
    chunks: List[Dict[str, Any]]
    status: ProcessingStatus
    processing_time: float
    error_message: Optional[str] = None
    ocr_confidence: Optional[float] = None
    performance_metrics: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            'document_id': self.document_id,
            'file_name': self.metadata.file_name,
            'file_size': self.metadata.file_size,
            'document_type': self.metadata.document_type.value,
            'page_count': self.metadata.page_count,
            'word_count': self.metadata.word_count,
            'text_length': len(self.extracted_text),
            'chunk_count': len(self.chunks),
            'status': self.status.value,
            'processing_time': self.processing_time,
            'error_message': self.error_message,
            'ocr_confidence': self.ocr_confidence,
            'performance_metrics': self.performance_metrics
        }


class DocumentProcessor:
    """Enhanced document processor with local OCR and security"""

    def __init__(self, config: ProcessingConfig = None):
        self.config = config or ProcessingConfig()
        self._validate_dependencies()

        # Create temp directory for processing
        self.temp_dir = Path(tempfile.mkdtemp(prefix="bear_ai_docs_"))

        logger.info(f"Document processor initialized with temp dir: {self.temp_dir}")

    def _validate_dependencies(self):
        """Validate required dependencies"""
        if self.config.ocr_enabled and not OCR_AVAILABLE:
            logger.warning("OCR requested but pytesseract/PIL not available")
            self.config.ocr_enabled = False

        if not PDF_AVAILABLE:
            logger.warning("PDF processing not available (PyMuPDF missing)")

        if not DOCX_AVAILABLE:
            logger.warning("DOCX processing not available (python-docx missing)")

    async def process_document(
        self,
        file_path: Union[str, Path],
        progress_callback: Optional[Callable[[float], None]] = None
    ) -> ProcessingResult:
        """Process a single document"""
        file_path = Path(file_path)
        document_id = str(uuid.uuid4())
        start_time = time.time()

        try:
            # Create metadata
            metadata = self._create_metadata(file_path)

            if progress_callback:
                progress_callback(0.1)

            # Extract text based on document type
            extracted_text, ocr_confidence = await self._extract_text(
                file_path, metadata, progress_callback
            )

            if progress_callback:
                progress_callback(0.6)

            # Clean and preprocess text
            cleaned_text = self._clean_text(extracted_text) if self.config.clean_text else extracted_text

            if progress_callback:
                progress_callback(0.8)

            # Create chunks
            chunks = self._create_chunks(cleaned_text, metadata)

            # Update metadata with extracted info
            metadata.word_count = len(cleaned_text.split())

            processing_time = time.time() - start_time

            result = ProcessingResult(
                document_id=document_id,
                metadata=metadata,
                extracted_text=extracted_text,
                cleaned_text=cleaned_text,
                chunks=chunks,
                status=ProcessingStatus.COMPLETED,
                processing_time=processing_time,
                ocr_confidence=ocr_confidence,
                performance_metrics={
                    'characters_per_second': len(extracted_text) / processing_time if processing_time > 0 else 0,
                    'words_per_second': len(cleaned_text.split()) / processing_time if processing_time > 0 else 0,
                    'chunk_count': len(chunks)
                }
            )

            if progress_callback:
                progress_callback(1.0)

            return result

        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"Document processing failed: {e}")

            return ProcessingResult(
                document_id=document_id,
                metadata=self._create_metadata(file_path) if 'metadata' not in locals() else metadata,
                extracted_text="",
                cleaned_text="",
                chunks=[],
                status=ProcessingStatus.FAILED,
                processing_time=processing_time,
                error_message=str(e)
            )

        finally:
            # Cleanup temp files
            if self.config.temp_file_cleanup:
                self._cleanup_temp_files()

    def _create_metadata(self, file_path: Path) -> DocumentMetadata:
        """Create document metadata"""
        try:
            stat = file_path.stat()
            mime_type, _ = mimetypes.guess_type(str(file_path))

            # Detect document type from extension
            doc_type = self._detect_document_type(file_path, mime_type)

            return DocumentMetadata(
                file_path=file_path,
                file_name=file_path.name,
                file_size=stat.st_size,
                mime_type=mime_type or "application/octet-stream",
                document_type=doc_type,
                creation_time=stat.st_ctime,
                modification_time=stat.st_mtime
            )
        except Exception as e:
            logger.error(f"Failed to create metadata for {file_path}: {e}")
            return DocumentMetadata(
                file_path=file_path,
                file_name=file_path.name,
                file_size=0,
                mime_type="unknown",
                document_type=DocumentType.UNKNOWN
            )

    def _detect_document_type(self, file_path: Path, mime_type: Optional[str]) -> DocumentType:
        """Detect document type from file extension and MIME type"""
        extension = file_path.suffix.lower()

        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOC,
            '.txt': DocumentType.TXT,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.rtf': DocumentType.RTF,
            '.xlsx': DocumentType.XLSX,
            '.xls': DocumentType.XLS,
            '.csv': DocumentType.CSV,
            '.png': DocumentType.IMAGE,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
            '.tiff': DocumentType.IMAGE,
            '.bmp': DocumentType.IMAGE
        }

        return type_mapping.get(extension, DocumentType.UNKNOWN)

    async def _extract_text(
        self,
        file_path: Path,
        metadata: DocumentMetadata,
        progress_callback: Optional[Callable[[float], None]] = None
    ) -> Tuple[str, Optional[float]]:
        """Extract text from document based on type"""

        if metadata.document_type == DocumentType.PDF:
            return await self._extract_pdf_text(file_path, metadata, progress_callback)
        elif metadata.document_type == DocumentType.DOCX:
            return await self._extract_docx_text(file_path, metadata)
        elif metadata.document_type == DocumentType.TXT:
            return await self._extract_txt_text(file_path, metadata)
        elif metadata.document_type == DocumentType.HTML:
            return await self._extract_html_text(file_path, metadata)
        elif metadata.document_type == DocumentType.IMAGE:
            return await self._extract_image_text(file_path, metadata)
        else:
            # Fallback to text extraction
            return await self._extract_txt_text(file_path, metadata)

    async def _extract_pdf_text(
        self,
        file_path: Path,
        metadata: DocumentMetadata,
        progress_callback: Optional[Callable[[float], None]] = None
    ) -> Tuple[str, Optional[float]]:
        """Extract text from PDF with OCR fallback"""
        if not PDF_AVAILABLE:
            raise RuntimeError("PDF processing not available")

        try:
            doc = fitz.open(str(file_path))
            metadata.page_count = doc.page_count

            text_pages = []
            ocr_pages = []
            total_confidence = 0.0
            confidence_count = 0

            for page_num in range(doc.page_count):
                if progress_callback:
                    progress = 0.2 + (page_num / doc.page_count) * 0.4
                    progress_callback(progress)

                page = doc[page_num]
                text = page.get_text().strip()

                if text:
                    text_pages.append(text)
                elif self.config.ocr_enabled:
                    # Extract as image for OCR
                    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                    img_data = pix.tobytes("png")

                    # Save temp image
                    temp_img_path = self.temp_dir / f"page_{page_num}.png"
                    with open(temp_img_path, 'wb') as f:
                        f.write(img_data)

                    ocr_text, confidence = await self._ocr_image(temp_img_path)
                    text_pages.append(ocr_text)

                    if confidence is not None:
                        total_confidence += confidence
                        confidence_count += 1
                else:
                    text_pages.append("")

            doc.close()

            full_text = '\n\n'.join(text_pages)
            avg_confidence = total_confidence / confidence_count if confidence_count > 0 else None

            return full_text, avg_confidence

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise

    async def _extract_docx_text(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, Optional[float]]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("DOCX processing not available")

        try:
            doc = DocxDocument(str(file_path))

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract table content if configured
            if self.config.extract_clauses:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        paragraphs.append(table_text)

            full_text = '\n\n'.join(paragraphs)
            return full_text, None

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    def _extract_table_text(self, table) -> str:
        """Extract text from DOCX table"""
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                table_rows.append(' | '.join(cells))
        return '\n'.join(table_rows)

    async def _extract_txt_text(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, Optional[float]]:
        """Extract text from plain text file"""
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    metadata.encoding = encoding
                    return text, None
                except UnicodeDecodeError:
                    continue

            # If all fail, use error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            metadata.encoding = 'utf-8'
            return text, None

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            raise

    async def _extract_html_text(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, Optional[float]]:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            if HTML_AVAILABLE:
                soup = BeautifulSoup(html_content, 'html.parser')

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                text = soup.get_text()

                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                clean_text = ' '.join(chunk for chunk in chunks if chunk)

                return clean_text, None
            else:
                # Simple HTML tag removal
                import re
                text = re.sub(r'<[^>]+>', '', html_content)
                return text, None

        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            raise

    async def _extract_image_text(
        self,
        file_path: Path,
        metadata: DocumentMetadata
    ) -> Tuple[str, Optional[float]]:
        """Extract text from image using OCR"""
        if not self.config.ocr_enabled or not OCR_AVAILABLE:
            return "", None

        return await self._ocr_image(file_path)

    async def _ocr_image(self, image_path: Path) -> Tuple[str, Optional[float]]:
        """Perform OCR on image"""
        try:
            # Load and preprocess image
            image = Image.open(image_path)

            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')

            # Enhance image for better OCR
            image = ImageEnhance.Contrast(image).enhance(1.5)
            image = ImageEnhance.Sharpness(image).enhance(1.2)
            image = image.filter(ImageFilter.SHARPEN)

            # Perform OCR with confidence
            ocr_config = f'--psm {self.config.ocr_psm} --dpi {self.config.ocr_dpi} -c tessedit_create_tsv=1'

            # Get text and confidence
            text = pytesseract.image_to_string(
                image,
                lang=self.config.ocr_language,
                config=ocr_config
            )

            # Get confidence data
            try:
                data = pytesseract.image_to_data(
                    image,
                    lang=self.config.ocr_language,
                    config=ocr_config,
                    output_type=pytesseract.Output.DICT
                )

                confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            except:
                avg_confidence = None

            return text.strip(), avg_confidence

        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return "", None

    def _clean_text(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        if not text:
            return ""

        import re

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove special characters that don't add value
        text = re.sub(r'[^\w\s\.\,\!\?\:\;\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>]', '', text)

        # Fix common OCR errors
        ocr_corrections = {
            r'\b1\b': 'I',  # Common OCR mistake
            r'\b0\b': 'O',  # Common OCR mistake
            r'([a-z])\1{3,}': r'\1',  # Remove repeated characters
        }

        for pattern, replacement in ocr_corrections.items():
            text = re.sub(pattern, replacement, text)

        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        text = text.strip()

        return text

    def _create_chunks(self, text: str, metadata: DocumentMetadata) -> List[Dict[str, Any]]:
        """Create text chunks with metadata"""
        if not text:
            return []

        # Smart chunking based on document structure
        chunks = []
        sentences = self._split_sentences(text)

        current_chunk = []
        current_size = 0
        chunk_index = 0

        for sentence in sentences:
            sentence_words = len(sentence.split())

            # Check if adding this sentence would exceed chunk size
            if current_size + sentence_words > self.config.chunk_size and current_chunk:
                # Create chunk
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    'id': str(uuid.uuid4()),
                    'text': chunk_text,
                    'chunk_index': chunk_index,
                    'word_count': current_size,
                    'document_id': metadata.checksum,
                    'document_type': metadata.document_type.value,
                    'file_name': metadata.file_name
                })

                # Start new chunk with overlap
                overlap_sentences = max(0, len(current_chunk) - self.config.chunk_overlap)
                current_chunk = current_chunk[overlap_sentences:] + [sentence]
                current_size = sum(len(s.split()) for s in current_chunk)
                chunk_index += 1
            else:
                current_chunk.append(sentence)
                current_size += sentence_words

        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'id': str(uuid.uuid4()),
                'text': chunk_text,
                'chunk_index': chunk_index,
                'word_count': current_size,
                'document_id': metadata.checksum,
                'document_type': metadata.document_type.value,
                'file_name': metadata.file_name
            })

        return chunks

    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences with legal document awareness"""
        import re

        # Legal abbreviations that shouldn't trigger sentence breaks
        legal_abbrevs = [
            'Inc', 'Corp', 'LLC', 'Ltd', 'Co', 'Esq', 'Jr', 'Sr',
            'v', 'vs', 'No', 'Sec', 'Art', 'Ch', 'Para', 'Subpara',
            'U.S', 'U.K', 'E.U', 'N.Y', 'Cal', 'Tex', 'Fla'
        ]

        # Protect abbreviations
        protected_text = text
        for abbrev in legal_abbrevs:
            protected_text = protected_text.replace(f'{abbrev}.', f'{abbrev}<!PERIOD!>')

        # Split on sentence boundaries
        sentences = re.split(r'(?<=[.!?])\s+', protected_text)

        # Restore periods
        sentences = [s.replace('<!PERIOD!>', '.').strip() for s in sentences if s.strip()]

        return sentences

    def _cleanup_temp_files(self):
        """Clean up temporary files"""
        try:
            import shutil
            if self.temp_dir.exists():
                shutil.rmtree(self.temp_dir)
        except Exception as e:
            logger.warning(f"Failed to cleanup temp files: {e}")

    def __del__(self):
        """Cleanup on destruction"""
        self._cleanup_temp_files()


# Factory function for easy creation
def create_document_processor(
    ocr_enabled: bool = True,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    clean_text: bool = True,
    extract_clauses: bool = True
) -> DocumentProcessor:
    """Create document processor with specified configuration"""
    config = ProcessingConfig(
        ocr_enabled=ocr_enabled,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        clean_text=clean_text,
        extract_clauses=extract_clauses
    )
    return DocumentProcessor(config)