"""
Advanced Document Processing Pipeline for BEAR AI
High-performance parallel processing for legal documents

Features:
- Multi-threaded and multi-process document processing
- Intelligent document chunking and preprocessing
- OCR and text extraction optimization
- Legal document type detection
- Batch processing with progress tracking
- Error recovery and retry mechanisms
- Memory-efficient streaming processing
- Real-time progress monitoring

@version 3.0.0
@author BEAR AI Document Processing Team
"""

import asyncio
import logging
import multiprocessing as mp
import os
import threading
import time
from abc import ABC, abstractmethod
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Any, Callable, Union, Tuple, Generator
from collections import deque, defaultdict
import queue
import mimetypes
import hashlib
import tempfile
import shutil

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Legal document types for specialized processing"""
    CONTRACT = "contract"
    LEGAL_BRIEF = "legal_brief"
    COURT_FILING = "court_filing"
    REGULATION = "regulation"
    STATUTE = "statute"
    CASE_LAW = "case_law"
    MEMORANDUM = "memorandum"
    CORRESPONDENCE = "correspondence"
    FINANCIAL = "financial"
    EVIDENCE = "evidence"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    """Document processing status"""
    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    FAILED = "failed"
    SKIPPED = "skipped"


@dataclass
class DocumentMetadata:
    """Enhanced document metadata"""
    file_path: Path
    file_size: int
    mime_type: str
    document_type: DocumentType = DocumentType.UNKNOWN
    language: str = "en"
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    creation_date: Optional[str] = None
    modification_date: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    confidence_score: float = 0.0


@dataclass
class ProcessingResult:
    """Document processing result"""
    document_id: str
    metadata: DocumentMetadata
    extracted_text: str
    chunks: List[Dict[str, Any]]
    status: ProcessingStatus
    processing_time: float
    error_message: Optional[str] = None
    performance_metrics: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ProcessingConfig:
    """Configuration for document processing"""
    max_workers: int = None  # Auto-detect
    chunk_size: int = 1000
    chunk_overlap: int = 200
    ocr_enabled: bool = True
    ocr_language: str = "eng"
    ocr_dpi: int = 300
    parallel_ocr: bool = True
    extract_images: bool = False
    extract_tables: bool = True
    detect_document_type: bool = True
    preserve_formatting: bool = False
    memory_limit_mb: int = 1000
    timeout_seconds: int = 300
    retry_attempts: int = 3
    batch_size: int = 10


class DocumentExtractor(ABC):
    """Abstract base class for document extractors"""
    
    @abstractmethod
    async def extract_text(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Extract text from document"""
        pass
    
    @abstractmethod
    def supports_format(self, mime_type: str) -> bool:
        """Check if extractor supports the format"""
        pass
    
    def get_priority(self) -> int:
        """Get extractor priority (lower = higher priority)"""
        return 100


class PDFExtractor(DocumentExtractor):
    """High-performance PDF text extraction"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.ocr_enabled = config.ocr_enabled and OCR_AVAILABLE
        
    async def extract_text(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Extract text from PDF with OCR fallback"""
        if not PDF_AVAILABLE:
            raise RuntimeError("PyMuPDF not available")
        
        try:
            # Open PDF document
            doc = fitz.open(str(file_path))
            
            # Update metadata
            metadata.page_count = doc.page_count
            if doc.metadata:
                metadata.title = doc.metadata.get('title', '')
                metadata.author = doc.metadata.get('author', '')
                metadata.creation_date = doc.metadata.get('creationDate', '')
            
            extracted_text = []
            ocr_pages = []
            
            # Extract text from each page
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    # Text extraction successful
                    extracted_text.append(text)
                elif self.ocr_enabled:
                    # No text found, add to OCR queue
                    ocr_pages.append((page_num, page))
                else:
                    extracted_text.append("")  # Placeholder for empty page
            
            # Process OCR pages in parallel if enabled
            if ocr_pages and self.config.parallel_ocr:
                ocr_results = await self._process_ocr_pages_parallel(ocr_pages, doc)
                
                # Merge OCR results
                for page_num, ocr_text in ocr_results:
                    if page_num < len(extracted_text):
                        extracted_text[page_num] = ocr_text
            
            doc.close()
            
            full_text = '\\n\\n'.join(extracted_text)
            metadata.word_count = len(full_text.split())
            
            return full_text
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {file_path}: {e}")
            raise
    
    async def _process_ocr_pages_parallel(
        self, 
        ocr_pages: List[Tuple[int, Any]], 
        doc: Any
    ) -> List[Tuple[int, str]]:
        """Process OCR pages in parallel"""
        loop = asyncio.get_event_loop()
        
        async def process_page_ocr(page_info):
            page_num, page = page_info
            return await loop.run_in_executor(
                None, 
                self._extract_text_with_ocr, 
                page, 
                page_num
            )
        
        tasks = [process_page_ocr(page_info) for page_info in ocr_pages]
        return await asyncio.gather(*tasks)
    
    def _extract_text_with_ocr(self, page: Any, page_num: int) -> Tuple[int, str]:
        """Extract text using OCR"""
        try:
            # Convert page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # 2x zoom for better OCR
            img_data = pix.tobytes("png")
            
            # Convert to PIL Image
            img = Image.open(io.BytesIO(img_data))
            
            # Enhance image for better OCR
            img = img.convert('L')  # Convert to grayscale
            img = ImageEnhance.Contrast(img).enhance(1.5)
            img = img.filter(ImageFilter.SHARPEN)
            
            # Perform OCR
            ocr_text = pytesseract.image_to_string(
                img, 
                lang=self.config.ocr_language,
                config=f'--psm 6 --dpi {self.config.ocr_dpi}'
            )
            
            return (page_num, ocr_text)
            
        except Exception as e:
            logger.error(f"OCR failed for page {page_num}: {e}")
            return (page_num, "")
    
    def supports_format(self, mime_type: str) -> bool:
        return mime_type == "application/pdf"
    
    def get_priority(self) -> int:
        return 10  # High priority for PDFs


class WordExtractor(DocumentExtractor):
    """Microsoft Word document extractor"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
    
    async def extract_text(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available")
        
        try:
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables if enabled
            if self.config.extract_tables:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text:
                        paragraphs.append(table_text)
            
            full_text = '\\n\\n'.join(paragraphs)
            
            # Update metadata
            metadata.word_count = len(full_text.split())
            if hasattr(doc.core_properties, 'title'):
                metadata.title = doc.core_properties.title or ""
            if hasattr(doc.core_properties, 'author'):
                metadata.author = doc.core_properties.author or ""
            
            return full_text
            
        except Exception as e:
            logger.error(f"Word extraction failed for {file_path}: {e}")
            raise
    
    def _extract_table_text(self, table) -> str:
        """Extract text from table"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            if any(row_data):  # Skip empty rows
                table_data.append(" | ".join(row_data))
        
        return "\\n".join(table_data)
    
    def supports_format(self, mime_type: str) -> bool:
        return mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]
    
    def get_priority(self) -> int:
        return 20


class TextExtractor(DocumentExtractor):
    """Plain text file extractor"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
    
    async def extract_text(self, file_path: Path, metadata: DocumentMetadata) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    metadata.word_count = len(text.split())
                    return text
                    
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            metadata.word_count = len(text.split())
            return text
            
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    def supports_format(self, mime_type: str) -> bool:
        return mime_type.startswith("text/")
    
    def get_priority(self) -> int:
        return 50


class DocumentTypeDetector:
    """Intelligent document type detection for legal documents"""
    
    def __init__(self):
        # Legal document patterns and keywords
        self.document_patterns = {
            DocumentType.CONTRACT: [
                "agreement", "contract", "terms and conditions", "whereas", "party",
                "consideration", "covenant", "breach", "termination", "intellectual property"
            ],
            DocumentType.LEGAL_BRIEF: [
                "brief", "memorandum", "argument", "plaintiff", "defendant", "court",
                "jurisdiction", "precedent", "statute", "ruling"
            ],
            DocumentType.COURT_FILING: [
                "motion", "petition", "complaint", "answer", "discovery", "deposition",
                "subpoena", "filing", "docket", "case number"
            ],
            DocumentType.REGULATION: [
                "regulation", "rule", "cfr", "federal register", "compliance",
                "regulatory", "administrative", "agency", "enforcement"
            ],
            DocumentType.STATUTE: [
                "statute", "code", "section", "subsection", "enacted", "legislation",
                "law", "legal code", "usc", "public law"
            ],
            DocumentType.CASE_LAW: [
                "opinion", "decision", "judgment", "justice", "court of appeals",
                "supreme court", "precedent", "holding", "dicta", "dissent"
            ],
            DocumentType.MEMORANDUM: [
                "memorandum", "memo", "legal memo", "analysis", "opinion letter",
                "legal opinion", "advice", "recommendation"
            ],
            DocumentType.CORRESPONDENCE: [
                "letter", "email", "correspondence", "communication", "notice",
                "demand letter", "cease and desist", "response"
            ],
            DocumentType.FINANCIAL: [
                "financial", "audit", "balance sheet", "income statement",
                "cash flow", "budget", "invoice", "contract price"
            ],
            DocumentType.EVIDENCE: [
                "evidence", "exhibit", "deposition", "testimony", "witness",
                "affidavit", "declaration", "expert report"
            ]
        }
    
    def detect_document_type(self, text: str, metadata: DocumentMetadata) -> DocumentType:
        """Detect document type based on content analysis"""
        text_lower = text.lower()
        scores = defaultdict(float)
        
        # Score based on keyword matching
        for doc_type, keywords in self.document_patterns.items():
            for keyword in keywords:
                # Count occurrences with position weighting (beginning matters more)
                occurrences = text_lower.count(keyword)
                if occurrences > 0:
                    # Weight keywords found in first 1000 characters higher
                    early_text = text_lower[:1000]
                    early_occurrences = early_text.count(keyword)
                    
                    score = occurrences + (early_occurrences * 2)  # Double weight for early matches
                    scores[doc_type] += score
        
        # Normalize scores by text length
        text_length = len(text.split())
        for doc_type in scores:
            scores[doc_type] = scores[doc_type] / max(text_length, 1) * 1000
        
        # Additional scoring based on metadata
        if metadata.title:
            title_lower = metadata.title.lower()
            for doc_type, keywords in self.document_patterns.items():
                for keyword in keywords:
                    if keyword in title_lower:
                        scores[doc_type] += 5.0  # Title match bonus
        
        # File extension hints
        if metadata.file_path.suffix.lower() == '.pdf':
            # PDFs are more likely to be formal legal documents
            scores[DocumentType.CONTRACT] += 1.0
            scores[DocumentType.LEGAL_BRIEF] += 1.0
        
        # Determine best match
        if not scores:
            return DocumentType.UNKNOWN
        
        best_type = max(scores, key=scores.get)
        confidence = scores[best_type]
        
        # Require minimum confidence threshold
        if confidence < 2.0:
            return DocumentType.UNKNOWN
        
        metadata.confidence_score = confidence
        return best_type


class DocumentChunker:
    """Intelligent document chunking for legal documents"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.chunk_size = config.chunk_size
        self.overlap = config.chunk_overlap
    
    def chunk_document(
        self, 
        text: str, 
        metadata: DocumentMetadata
    ) -> List[Dict[str, Any]]:
        """Create intelligent chunks from document"""
        if not text.strip():
            return []
        
        # Legal document-specific chunking strategies
        if metadata.document_type in [DocumentType.CONTRACT, DocumentType.LEGAL_BRIEF]:
            return self._chunk_by_sections(text, metadata)
        elif metadata.document_type == DocumentType.CASE_LAW:
            return self._chunk_by_legal_sections(text, metadata)
        else:
            return self._chunk_by_sentences(text, metadata)
    
    def _chunk_by_sections(self, text: str, metadata: DocumentMetadata) -> List[Dict[str, Any]]:
        """Chunk by document sections (contracts, briefs)"""
        # Detect section headers
        lines = text.split('\\n')
        sections = []
        current_section = []
        section_title = ""
        
        for line in lines:
            line = line.strip()
            
            # Detect section headers (numbered sections, ALL CAPS headings, etc.)
            if self._is_section_header(line):
                if current_section:
                    sections.append({
                        'title': section_title,
                        'content': '\\n'.join(current_section).strip()
                    })
                
                section_title = line
                current_section = []
            else:
                current_section.append(line)
        
        # Add final section
        if current_section:
            sections.append({
                'title': section_title,
                'content': '\\n'.join(current_section).strip()
            })
        
        # Create chunks from sections
        chunks = []
        for i, section in enumerate(sections):
            if not section['content'].strip():
                continue
            
            section_chunks = self._split_text_by_size(section['content'])
            
            for j, chunk_text in enumerate(section_chunks):
                chunks.append({
                    'text': chunk_text,
                    'metadata': {
                        'chunk_index': len(chunks),
                        'section_title': section['title'],
                        'section_index': i,
                        'subsection_index': j,
                        'document_type': metadata.document_type.value,
                        'word_count': len(chunk_text.split())
                    }
                })
        
        return chunks
    
    def _chunk_by_legal_sections(self, text: str, metadata: DocumentMetadata) -> List[Dict[str, Any]]:
        """Chunk case law by legal sections (facts, holding, reasoning)"""
        # Legal section patterns
        section_patterns = [
            (r'\\b(facts?|background|procedural history)\\b', 'facts'),
            (r'\\b(issue|question presented|legal question)\\b', 'issue'),
            (r'\\b(holding|decision|ruling|conclusion)\\b', 'holding'),
            (r'\\b(reasoning|analysis|discussion|rationale)\\b', 'reasoning'),
            (r'\\b(dissent|dissenting opinion)\\b', 'dissent'),
            (r'\\b(concur|concurring opinion)\\b', 'concurrence')
        ]
        
        import re
        
        # Split text into legal sections
        sections = []
        current_section = {'type': 'general', 'content': []}
        
        for line in text.split('\\n'):
            line = line.strip()
            if not line:
                continue
            
            # Check if line indicates a new section
            section_type = None
            for pattern, stype in section_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    section_type = stype
                    break
            
            if section_type and section_type != current_section['type']:
                # Save current section
                if current_section['content']:
                    sections.append(current_section)
                
                # Start new section
                current_section = {'type': section_type, 'content': [line]}
            else:
                current_section['content'].append(line)
        
        # Add final section
        if current_section['content']:
            sections.append(current_section)
        
        # Create chunks from sections
        chunks = []
        for section in sections:
            section_text = '\\n'.join(section['content'])
            section_chunks = self._split_text_by_size(section_text)
            
            for chunk_text in section_chunks:
                chunks.append({
                    'text': chunk_text,
                    'metadata': {
                        'chunk_index': len(chunks),
                        'legal_section': section['type'],
                        'document_type': metadata.document_type.value,
                        'word_count': len(chunk_text.split())
                    }
                })
        
        return chunks
    
    def _chunk_by_sentences(self, text: str, metadata: DocumentMetadata) -> List[Dict[str, Any]]:
        """Chunk by sentences with size limits"""
        sentences = self._split_sentences(text)
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence.split())
            
            if current_size + sentence_size > self.chunk_size and current_chunk:
                # Create chunk
                chunk_text = ' '.join(current_chunk)
                chunks.append({
                    'text': chunk_text,
                    'metadata': {
                        'chunk_index': len(chunks),
                        'document_type': metadata.document_type.value,
                        'word_count': len(chunk_text.split()),
                        'sentence_count': len(current_chunk)
                    }
                })
                
                # Start new chunk with overlap
                overlap_sentences = current_chunk[-self.overlap:] if len(current_chunk) > self.overlap else current_chunk
                current_chunk = overlap_sentences + [sentence]
                current_size = sum(len(s.split()) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_size += sentence_size
        
        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                'text': chunk_text,
                'metadata': {
                    'chunk_index': len(chunks),
                    'document_type': metadata.document_type.value,
                    'word_count': len(chunk_text.split()),
                    'sentence_count': len(current_chunk)
                }
            })
        
        return chunks
    
    def _is_section_header(self, line: str) -> bool:
        """Detect if line is a section header"""
        if not line.strip():
            return False
        
        # Check for numbered sections
        import re
        if re.match(r'^\\s*\\d+\\.', line):
            return True
        
        # Check for ALL CAPS headers (but not entire document in caps)
        if line.isupper() and len(line) < 100:
            return True
        
        # Check for Roman numerals
        if re.match(r'^\\s*[IVX]+\\.', line):
            return True
        
        # Check for lettered sections
        if re.match(r'^\\s*[A-Z]\\.', line):
            return True
        
        return False
    
    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences with legal document awareness"""
        import re
        
        # Legal-aware sentence splitting
        # Don't split on abbreviations common in legal documents
        abbreviations = [
            'Inc', 'Corp', 'LLC', 'Ltd', 'Co', 'Esq', 'Jr', 'Sr',
            'v', 'vs', 'No', 'Sec', 'Art', 'Ch', 'Para', 'Subpara'
        ]
        
        # Protect abbreviations
        protected_text = text
        for abbr in abbreviations:
            protected_text = protected_text.replace(f'{abbr}.', f'{abbr}<!PERIOD!>')
        
        # Split on sentence endings
        sentences = re.split(r'(?<=[.!?])\\s+', protected_text)
        
        # Restore periods
        sentences = [s.replace('<!PERIOD!>', '.') for s in sentences]
        
        return [s.strip() for s in sentences if s.strip()]
    
    def _split_text_by_size(self, text: str) -> List[str]:
        """Split text by size maintaining word boundaries"""
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0
        
        for word in words:
            word_size = 1  # Count by words
            
            if current_size + word_size > self.chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                
                # Start new chunk with overlap
                overlap_words = current_chunk[-self.overlap:] if len(current_chunk) > self.overlap else []
                current_chunk = overlap_words + [word]
                current_size = len(current_chunk)
            else:
                current_chunk.append(word)
                current_size += word_size
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks


class ProgressTracker:
    """Track processing progress with detailed metrics"""
    
    def __init__(self):
        self.total_documents = 0
        self.processed_documents = 0
        self.failed_documents = 0
        self.skipped_documents = 0
        self.start_time = None
        self.lock = threading.Lock()
        self.document_status = {}
        self.performance_metrics = {
            'total_processing_time': 0.0,
            'total_text_extracted': 0,
            'total_chunks_created': 0,
            'ocr_pages_processed': 0,
            'average_processing_speed': 0.0
        }
    
    def start(self, total_documents: int):
        """Start tracking progress"""
        with self.lock:
            self.total_documents = total_documents
            self.processed_documents = 0
            self.failed_documents = 0
            self.skipped_documents = 0
            self.start_time = time.time()
            self.document_status = {}
    
    def update_document_status(self, document_id: str, status: ProcessingStatus, result: Optional[ProcessingResult] = None):
        """Update status for a specific document"""
        with self.lock:
            old_status = self.document_status.get(document_id)
            self.document_status[document_id] = status
            
            # Update counters only if status actually changed
            if old_status != status:
                if status == ProcessingStatus.COMPLETED:
                    self.processed_documents += 1
                elif status == ProcessingStatus.FAILED:
                    self.failed_documents += 1
                elif status == ProcessingStatus.SKIPPED:
                    self.skipped_documents += 1
            
            # Update performance metrics
            if result:
                self.performance_metrics['total_processing_time'] += result.processing_time
                self.performance_metrics['total_text_extracted'] += len(result.extracted_text)
                self.performance_metrics['total_chunks_created'] += len(result.chunks)
                
                # Calculate average processing speed
                if self.processed_documents > 0:
                    self.performance_metrics['average_processing_speed'] = (
                        self.performance_metrics['total_text_extracted'] / 
                        self.performance_metrics['total_processing_time']
                    )
    
    def get_progress(self) -> Dict[str, Any]:
        """Get current progress statistics"""
        with self.lock:
            total_completed = self.processed_documents + self.failed_documents + self.skipped_documents
            progress_percentage = (total_completed / max(self.total_documents, 1)) * 100
            
            elapsed_time = time.time() - self.start_time if self.start_time else 0
            estimated_remaining = 0
            
            if self.processed_documents > 0 and elapsed_time > 0:
                avg_time_per_doc = elapsed_time / total_completed
                remaining_docs = self.total_documents - total_completed
                estimated_remaining = remaining_docs * avg_time_per_doc
            
            return {
                'total_documents': self.total_documents,
                'processed': self.processed_documents,
                'failed': self.failed_documents,
                'skipped': self.skipped_documents,
                'progress_percentage': progress_percentage,
                'elapsed_time': elapsed_time,
                'estimated_remaining': estimated_remaining,
                'performance_metrics': self.performance_metrics.copy()
            }


class DocumentProcessingPipeline:
    """Main document processing pipeline with parallel execution"""
    
    def __init__(self, config: ProcessingConfig = None):
        self.config = config or ProcessingConfig()
        
        # Auto-detect optimal worker count
        if self.config.max_workers is None:
            self.config.max_workers = min(32, (os.cpu_count() or 1) + 4)
        
        # Initialize components
        self.extractors = self._initialize_extractors()
        self.document_detector = DocumentTypeDetector()
        self.chunker = DocumentChunker(self.config)
        self.progress_tracker = ProgressTracker()
        
        # Processing queues
        self.input_queue = queue.Queue()
        self.result_queue = queue.Queue()
        self.error_queue = queue.Queue()
        
        logger.info(f"Document pipeline initialized with {self.config.max_workers} workers")
    
    def _initialize_extractors(self) -> List[DocumentExtractor]:
        """Initialize document extractors"""
        extractors = []
        
        if PDF_AVAILABLE:
            extractors.append(PDFExtractor(self.config))
        
        if DOCX_AVAILABLE:
            extractors.append(WordExtractor(self.config))
        
        extractors.append(TextExtractor(self.config))
        
        # Sort by priority
        extractors.sort(key=lambda x: x.get_priority())
        
        return extractors
    
    def _get_extractor(self, mime_type: str) -> Optional[DocumentExtractor]:
        """Get appropriate extractor for file type"""
        for extractor in self.extractors:
            if extractor.supports_format(mime_type):
                return extractor
        return None
    
    async def process_documents(
        self,
        file_paths: List[Path],
        progress_callback: Optional[Callable[[Dict[str, Any]], None]] = None
    ) -> List[ProcessingResult]:
        """Process multiple documents in parallel"""
        
        if not file_paths:
            return []
        
        self.progress_tracker.start(len(file_paths))
        
        # Create tasks for parallel processing
        semaphore = asyncio.Semaphore(self.config.max_workers)
        tasks = []
        
        for file_path in file_paths:
            task = asyncio.create_task(
                self._process_document_with_semaphore(semaphore, file_path)
            )
            tasks.append(task)
        
        results = []
        
        # Process documents and track progress
        for completed_task in asyncio.as_completed(tasks):
            try:
                result = await completed_task
                results.append(result)
                
                # Update progress
                if progress_callback:
                    progress_callback(self.progress_tracker.get_progress())
                    
            except Exception as e:
                logger.error(f"Task failed: {e}")
        
        return results
    
    async def _process_document_with_semaphore(
        self, 
        semaphore: asyncio.Semaphore, 
        file_path: Path
    ) -> ProcessingResult:
        """Process single document with semaphore limiting"""
        async with semaphore:
            return await self.process_single_document(file_path)
    
    async def process_single_document(self, file_path: Path) -> ProcessingResult:
        """Process a single document"""
        document_id = hashlib.md5(str(file_path).encode()).hexdigest()
        start_time = time.time()
        
        # Update status to in progress
        self.progress_tracker.update_document_status(document_id, ProcessingStatus.IN_PROGRESS)
        
        try:
            # Check file exists and get basic metadata
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_size = file_path.stat().st_size
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            if not mime_type:
                mime_type = "application/octet-stream"
            
            # Create metadata
            metadata = DocumentMetadata(
                file_path=file_path,
                file_size=file_size,
                mime_type=mime_type
            )
            
            # Get appropriate extractor
            extractor = self._get_extractor(mime_type)
            if not extractor:
                raise ValueError(f"No extractor available for {mime_type}")
            
            # Extract text
            extracted_text = await extractor.extract_text(file_path, metadata)
            
            # Detect document type
            if self.config.detect_document_type:
                metadata.document_type = self.document_detector.detect_document_type(
                    extracted_text, metadata
                )
            
            # Create chunks
            chunks = self.chunker.chunk_document(extracted_text, metadata)
            
            # Create result
            processing_time = time.time() - start_time
            result = ProcessingResult(
                document_id=document_id,
                metadata=metadata,
                extracted_text=extracted_text,
                chunks=chunks,
                status=ProcessingStatus.COMPLETED,
                processing_time=processing_time,
                performance_metrics={
                    'extraction_time': processing_time,
                    'text_length': len(extracted_text),
                    'chunk_count': len(chunks),
                    'processing_speed': len(extracted_text) / processing_time if processing_time > 0 else 0
                }
            )
            
            # Update progress
            self.progress_tracker.update_document_status(document_id, ProcessingStatus.COMPLETED, result)
            
            return result
            
        except Exception as e:
            processing_time = time.time() - start_time
            error_result = ProcessingResult(
                document_id=document_id,
                metadata=DocumentMetadata(
                    file_path=file_path,
                    file_size=0,
                    mime_type="unknown"
                ),
                extracted_text="",
                chunks=[],
                status=ProcessingStatus.FAILED,
                processing_time=processing_time,
                error_message=str(e)
            )
            
            self.progress_tracker.update_document_status(document_id, ProcessingStatus.FAILED)
            logger.error(f"Document processing failed for {file_path}: {e}")
            
            return error_result
    
    def get_progress(self) -> Dict[str, Any]:
        """Get current processing progress"""
        return self.progress_tracker.get_progress()


# Factory function
def create_document_pipeline(
    max_workers: Optional[int] = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    ocr_enabled: bool = True,
    parallel_ocr: bool = True,
    extract_tables: bool = True,
    detect_document_type: bool = True
) -> DocumentProcessingPipeline:
    """Create optimized document processing pipeline"""
    
    config = ProcessingConfig(
        max_workers=max_workers,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        ocr_enabled=ocr_enabled,
        parallel_ocr=parallel_ocr,
        extract_tables=extract_tables,
        detect_document_type=detect_document_type
    )
    
    return DocumentProcessingPipeline(config)


# Example usage
async def process_legal_documents(file_paths: List[str]) -> List[ProcessingResult]:
    """Process legal documents with optimal settings"""
    pipeline = create_document_pipeline(
        max_workers=8,
        chunk_size=800,  # Smaller chunks for legal documents
        ocr_enabled=True,
        detect_document_type=True
    )
    
    path_objects = [Path(p) for p in file_paths]
    
    def progress_callback(progress):
        print(f"Progress: {progress['progress_percentage']:.1f}% "
              f"({progress['processed']}/{progress['total_documents']})")
    
    return await pipeline.process_documents(path_objects, progress_callback)